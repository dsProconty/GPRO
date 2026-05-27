from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_table_from_rows(doc, rows):
    if not rows:
        return
    # Detect header separator row (---|---|---)
    data_rows = [r for r in rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)
    if not parsed:
        return
    max_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=max_cols)
    table.style = 'Table Grid'
    for i, row in enumerate(parsed):
        for j, cell_text in enumerate(row):
            if j < max_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                    cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'BDD7EE')
                    shd.set(qn('w:val'), 'clear')
                    cell._tc.tcPr.append(shd)
    doc.add_paragraph()

lines = open('/home/user/GPRO/DEPLOYMENT_AZURE.md').readlines()

i = 0
table_rows = []
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Collect table rows
    if line.startswith('|'):
        table_rows.append(line)
        i += 1
        continue
    else:
        if table_rows:
            add_table_from_rows(doc, table_rows)
            table_rows = []

    if line.startswith('# '):
        add_heading(doc, line[2:], 1)
    elif line.startswith('## '):
        add_heading(doc, line[3:], 2)
    elif line.startswith('### '):
        add_heading(doc, line[4:], 3)
    elif line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    elif line.startswith('```'):
        # collect code block
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        p = doc.add_paragraph('\n'.join(code_lines))
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    elif line.startswith('---'):
        doc.add_paragraph('─' * 60)
    elif line.startswith('**') and line.endswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(line.strip('*'))
        run.bold = True
    elif line.strip() == '':
        pass
    else:
        # Handle inline bold/italic
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            else:
                p.add_run(part)
    i += 1

if table_rows:
    add_table_from_rows(doc, table_rows)

doc.save('/home/user/GPRO/DEPLOYMENT_AZURE.docx')
print("Done")
